#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convert DOCX to Markdown and extract images
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import os
import re
from pathlib import Path
import json


def extract_images_and_convert_to_md(docx_path, output_md_path, images_dir):
    """
    将DOCX文件转换为Markdown，并提取图片
    """
    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_md_path), exist_ok=True)
    os.makedirs(images_dir, exist_ok=True)

    doc = Document(docx_path)

    markdown_content = []
    image_counter = 0
    image_map = {}  # 记录图片映射

    # 处理每个元素
    for para in doc.paragraphs:
        # 提取段落中的文本
        text = para.text.strip()

        if text:
            # 根据段落风格判断标题级别
            if para.style.name.startswith("Heading"):
                level = (
                    int(para.style.name.split(" ")[-1]) if " " in para.style.name else 1
                )
                markdown_content.append(f"{'#' * level} {text}\n")
            elif para.style.name == "List Bullet":
                markdown_content.append(f"- {text}\n")
            elif para.style.name == "List Number":
                markdown_content.append(f"1. {text}\n")
            else:
                markdown_content.append(f"{text}\n")

        # 提取段落中的图片
        for run in para.runs:
            for inline_shape in run._element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
            ):
                image_filename = extract_image_from_element(
                    doc, inline_shape, images_dir, image_counter
                )
                if image_filename:
                    image_counter += 1
                    rel_path = f"pic/{image_filename}"
                    markdown_content.append(f"![image]({rel_path})\n")
                    image_map[image_filename] = True

    # 处理表格中的图片
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        markdown_content.append(f"{text}\n")

                    for run in para.runs:
                        for inline_shape in run._element.findall(
                            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                        ):
                            image_filename = extract_image_from_element(
                                doc, inline_shape, images_dir, image_counter
                            )
                            if image_filename:
                                image_counter += 1
                                rel_path = f"pic/{image_filename}"
                                markdown_content.append(f"![image]({rel_path})\n")
                                image_map[image_filename] = True

    # 写入Markdown文件
    with open(output_md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(markdown_content))

    print(f"✅ Markdown文件已生成: {output_md_path}")
    print(f"✅ 提取了 {image_counter} 张图片到 {images_dir}")

    return image_map


def extract_image_from_element(doc, element, output_dir, index):
    """
    从DOCX中提取图片
    """
    try:
        # 获取图片关系信息
        blip = element.find(
            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
        )
        if blip is None:
            return None

        embed_id = blip.get(qn("r:embed"))
        if not embed_id:
            return None

        # 获取图片数据 使用正确的API
        try:
            image_part = doc.core_properties._element.getparent().getall_relparts()[
                embed_id
            ]
        except Exception:
            # 尝试另一种方式
            try:
                rel_parts = doc.part.rels
                if embed_id in rel_parts:
                    image_part = rel_parts[embed_id].target_part
                else:
                    return None
            except Exception:
                return None

        try:
            image_bytes = image_part.blob
        except Exception:
            return None

        # 提取扩展名
        content_type = image_part.content_type
        ext = content_type.split("/")[-1]
        if ext == "jpeg":
            ext = "jpg"

        # 生成文件名
        filename = f"history_{index}.{ext}"
        filepath = os.path.join(output_dir, filename)

        # 保存图片
        with open(filepath, "wb") as f:
            f.write(image_bytes)

        print(f"  📸 已保存图片: {filename}")
        return filename

    except Exception as e:
        print(f"  ❌ 提取图片失败: {e}")
        return None


if __name__ == "__main__":
    docx_file = "牛腩史书.docx"
    output_md = "docs/牛腩史书.md"
    images_dir = "public/pic"

    if os.path.exists(docx_file):
        image_map = extract_images_and_convert_to_md(docx_file, output_md, images_dir)
        print("\n处理完成！")
        print(f"提取的文件数: {len(image_map)}")
    else:
        print(f"❌ 文件不存在: {docx_file}")
